"""
Tools for crawling, AI calls, and document generation.
"""

from datetime import datetime
import io
import logging
import random
import re
import time
from typing import Any, Dict, Optional
from urllib.parse import urlparse

import anthropic
from bs4 import BeautifulSoup
from docx import Document
from fake_useragent import UserAgent
import openai
import requests

from config.settings import NEWS_SOURCES


class UniversalWebCrawlerTool:
    """Crawl article pages from supported news sources."""

    def __init__(self):
        self.session = requests.Session()
        self.ua = UserAgent()
        self.logger = logging.getLogger(__name__)
        self.session.headers.update(
            {
                "User-Agent": self.ua.random,
                "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
                "Accept-Language": "vi-VN,vi;q=0.9,en;q=0.8",
                "Accept-Encoding": "gzip, deflate",
                "Connection": "keep-alive",
            }
        )

    def fetch_page(self, url: str, retries: int = 3) -> Optional[BeautifulSoup]:
        """Fetch a page and handle simple cookie-based anti-bot gates."""
        for attempt in range(retries):
            try:
                time.sleep(random.uniform(0.5, 2.0))
                self.session.headers["User-Agent"] = self.ua.random
                response = self.session.get(url, timeout=30, allow_redirects=True)
                if response.status_code != 200:
                    self.logger.warning(f"HTTP {response.status_code}: {url}")
                    continue

                if self._is_cookie_gate_response(response.text):
                    if not self._apply_cookie_gate(response.text, url):
                        self.logger.warning(f"Cookie gate without usable token: {url}")
                        continue
                    response = self.session.get(url, timeout=30, allow_redirects=True)
                    if response.status_code != 200:
                        self.logger.warning(f"HTTP {response.status_code} after cookie retry: {url}")
                        continue

                if self._is_cookie_gate_response(response.text):
                    self.logger.warning(f"Still blocked by cookie gate: {url}")
                    continue

                self.logger.info(f"Crawl thanh cong: {url}")
                return BeautifulSoup(response.content, "html.parser")
            except Exception as exc:
                self.logger.error(f"Loi crawl (attempt {attempt + 1}): {exc}")
                if attempt < retries - 1:
                    time.sleep(2**attempt)
        return None

    def _is_cookie_gate_response(self, html: str) -> bool:
        if not html:
            return False
        lower_html = html.lower()
        return "document.cookie=" in lower_html and "window.location.reload" in lower_html

    def _apply_cookie_gate(self, html: str, url: str) -> bool:
        match = re.search(r'document\.cookie="([^=]+)=([^"]+)"', html)
        if not match:
            return False

        cookie_name, cookie_value = match.group(1), match.group(2)
        domain = urlparse(url).netloc.replace("www.", "")
        self.session.cookies.set(cookie_name, cookie_value, domain=domain, path="/")
        self.logger.info(f"Da nhan cookie gate cho {domain}, thu crawl lai")
        return True

    def auto_extract_content(self, soup: BeautifulSoup, url: str) -> Optional[Dict[str, str]]:
        """Auto extract article content from a page."""
        try:
            source = self._detect_source_from_url(url)
            title = self._extract_title(soup)
            content = self._extract_main_content(soup, source)
            publish_time = self._extract_publish_time(soup)

            if title and content and len(content) > 50:
                return {
                    "title": title,
                    "content": content,
                    "source": source,
                    "publish_time": publish_time,
                    "url": url,
                }
            return None
        except Exception as exc:
            self.logger.error(f"Loi auto extract tu {url}: {exc}")
            return None

    def _detect_source_from_url(self, url: str) -> str:
        url_lower = url.lower()
        source_map = {
            "nhandan.vn": "nhandan",
            "vtv.vn": "vtv",
            "qdnd.vn": "qdnd",
            "tuoitre.vn": "tuoitre",
            "vnexpress.net": "vnexpress",
            "dantri.com.vn": "dantri",
            "thanhnien.vn": "thanhnien",
            "laodong.vn": "laodong",
            "nld.com.vn": "nld",
            "cand.com.vn": "cand",
            "baophapluat.vn": "baophapluat",
            "vietnamnet.vn": "vietnamnet",
            "baomoi.com": "baomoi",
        }
        for domain, source in source_map.items():
            if domain in url_lower:
                return source

        domain = urlparse(url).netloc
        return domain.replace("www.", "").replace(".vn", "").replace(".com", "").replace(".net", "")

    def _extract_title(self, soup: BeautifulSoup) -> Optional[str]:
        title_selectors = [
            "h1.title-detail",
            "h1.detail-title",
            "h1.article-title",
            ".story-title h1",
            "h1",
            ".title-detail",
            ".detail-title",
            ".article-title",
            "title",
        ]
        for selector in title_selectors:
            try:
                for element in soup.select(selector):
                    text = element.get_text(" ", strip=True)
                    if text and 10 < len(text) < 300:
                        return self._clean_title(text)
            except Exception:
                pass
        return None

    def _clean_title(self, title: str) -> str:
        suffixes = [
            " - QDND",
            " | QDND",
            " - VTV",
            " | VTV",
            " - Nhan Dan",
            " | Nhan Dan",
            " - Tuoi Tre Online",
            " | Bao Dan tri",
            " | Dan tri",
            " | Thanh Nien",
            " | Lao Dong",
        ]
        for suffix in suffixes:
            title = title.replace(suffix, "")
        return re.sub(r"\s+", " ", title).strip()

    def _extract_main_content(self, soup: BeautifulSoup, source: str = None) -> Optional[str]:
        content_parts: list[str] = []

        if source and source in NEWS_SOURCES:
            selectors = NEWS_SOURCES[source].get("selectors", {}).get("content", "")
            for selector in [item.strip() for item in selectors.split(",") if item.strip()]:
                content_parts.extend(self._collect_text_from_selector(soup, selector, source))
                if len(content_parts) >= 8:
                    break

        if not content_parts:
            common_selectors = [
                "article p",
                ".article-content p",
                ".detail-content p",
                ".story-content p",
                ".story-body p",
                ".content-detail p",
                ".content-body p",
                "main p",
                "p",
            ]
            for selector in common_selectors:
                content_parts.extend(self._collect_text_from_selector(soup, selector, source))
                if len(content_parts) >= 8:
                    break

        if not content_parts:
            return None

        content_parts = self._deduplicate_content_parts(content_parts)
        raw_content = "\n".join(content_parts[:40])
        cleaned = ContentCleanerTool.clean_text(raw_content)
        return cleaned if cleaned and len(cleaned) > 50 else None

    def _collect_text_from_selector(
        self, soup: BeautifulSoup, selector: str, source: Optional[str]
    ) -> list[str]:
        parts = []
        try:
            for element in soup.select(selector):
                text = self._extract_clean_text_from_element(element, source)
                if text:
                    parts.append(text)
        except Exception:
            pass
        return parts

    def _extract_clean_text_from_element(self, element, source: Optional[str] = None) -> Optional[str]:
        if not element:
            return None

        text = element.get_text(" ", strip=True)
        text = re.sub(r"\s+", " ", text).strip()
        if not text or len(text) < 20:
            return None
        if self._is_noise_text(text, source):
            return None
        return text

    def _is_noise_text(self, text: str, source: Optional[str]) -> bool:
        lower = text.lower().strip()

        generic_noise = [
            "ảnh:",
            "video:",
            "xem thêm",
            "bình luận",
            "quảng cáo",
            "liên hệ",
            "theo dõi",
            "đăng nhập",
            "đăng ký",
            "chia sẻ",
            "facebook",
            "youtube",
            "zalo",
            "email",
            "hotline",
            "bạn đang đọc",
        ]
        if any(token in lower for token in generic_noise):
            return True

        if source == "tuoitre":
            if lower.startswith("ảnh ") or lower.startswith("ảnh:") or "tuổi trẻ online" in lower:
                return True
        elif source == "dantri":
            if lower.startswith("(dân trí)") or lower.startswith("ảnh:"):
                return True
        elif source == "thanhnien":
            if lower.startswith("(tno)") or lower.startswith("ảnh:") or "video đang được phát" in lower:
                return True
        elif source == "laodong":
            if lower.startswith("minh họa:") or lower.startswith("ảnh:") or "xem thêm video" in lower:
                return True

        if len(lower.split()) <= 3 and any(char.isdigit() for char in lower):
            return True

        return False

    def _deduplicate_content_parts(self, content_parts: list[str]) -> list[str]:
        cleaned_parts: list[str] = []
        seen = set()

        for part in content_parts:
            normalized = re.sub(r"\s+", " ", part).strip()
            if not normalized:
                continue
            key = normalized.lower()
            if key in seen:
                continue
            if any(key in existing.lower() or existing.lower() in key for existing in cleaned_parts):
                continue
            seen.add(key)
            cleaned_parts.append(normalized)

        return cleaned_parts

    def _extract_publish_time(self, soup: BeautifulSoup) -> Optional[str]:
        time_selectors = [
            "time",
            ".publish-time",
            ".date",
            ".date-time",
            ".detail-time",
            ".author-time",
            ".time-post",
            "[class*='time']",
            "[class*='date']",
            "meta[property='article:published_time']",
            "meta[name='pubdate']",
        ]
        for selector in time_selectors:
            try:
                for element in soup.select(selector):
                    datetime_attr = element.get("datetime") or element.get("content")
                    if datetime_attr:
                        return datetime_attr
                    text = element.get_text(" ", strip=True)
                    if text and len(text) > 5:
                        return text
            except Exception:
                pass
        return None

    def fetch_latest_articles(self, source: str, num_articles: int = 3) -> Optional[list]:
        source_lower = source.lower()
        if source_lower not in NEWS_SOURCES:
            self.logger.warning(f"Nguon {source} khong duoc ho tro")
            return None

        source_config = NEWS_SOURCES[source_lower]
        articles = []
        for category_url in source_config.get("categories", []):
            if len(articles) >= num_articles:
                break
            try:
                soup = self.fetch_page(category_url)
                if not soup:
                    continue
                article_links = self._extract_article_links(soup, source_config)
                for link in article_links:
                    if len(articles) >= num_articles:
                        break
                    full_url = (
                        link
                        if link.startswith("http")
                        else source_config["base_url"].rstrip("/") + "/" + link.lstrip("/")
                    )
                    page = self.fetch_page(full_url)
                    if not page:
                        continue
                    article_data = self.auto_extract_content(page, full_url)
                    if article_data and len(article_data.get("content", "")) > 100:
                        articles.append(
                            {
                                "url": full_url,
                                "title": article_data.get("title"),
                                "content": article_data.get("content"),
                                "source": source_lower,
                            }
                        )
            except Exception as exc:
                self.logger.error(f"Loi fetch tu {category_url}: {exc}")
        return articles if articles else None

    def _extract_article_links(self, soup: BeautifulSoup, source_config: Dict) -> list:
        article_links = []
        article_link_selector = source_config.get("selectors", {}).get("article_links", "a[href]")
        try:
            for element in soup.select(article_link_selector):
                href = element.get("href")
                if href:
                    article_links.append(href)
        except Exception as exc:
            self.logger.warning(f"Loi trich xuat links: {exc}")
        return article_links


class AIModelTool:
    """Interface to supported AI models."""

    def __init__(self, model: str, api_key: str):
        self.model = model
        self.api_key = api_key
        self.logger = logging.getLogger(__name__)

        if "gpt" in model.lower():
            openai.api_key = api_key
            self.client_type = "openai"
        elif "claude" in model.lower():
            self.client = anthropic.Anthropic(api_key=api_key)
            self.client_type = "anthropic"
        else:
            raise ValueError(f"Model {model} khong duoc ho tro")

    def generate_text(
        self,
        system_prompt: str,
        user_prompt: str,
        max_tokens: int = 2000,
        temperature: float = 0.7,
    ) -> Optional[str]:
        try:
            total_chars = len(system_prompt) + len(user_prompt)
            estimated_tokens = total_chars // 3
            max_input_tokens = 14000

            if estimated_tokens > max_input_tokens:
                self.logger.warning(f"Input qua dai: {estimated_tokens} tokens, dang cat ngan...")
                max_user_chars = max_input_tokens * 3 - len(system_prompt)
                if len(user_prompt) > max_user_chars:
                    user_prompt = (
                        user_prompt[:max_user_chars]
                        + "\n[...noi dung da duoc cat ngan de phu hop voi AI...]"
                    )

            if self.client_type == "openai":
                response = openai.ChatCompletion.create(
                    model=self.model,
                    messages=[
                        {"role": "system", "content": system_prompt},
                        {"role": "user", "content": user_prompt},
                    ],
                    max_tokens=max_tokens,
                    temperature=temperature,
                )
                return response.choices[0].message.content.strip()

            response = self.client.messages.create(
                model=self.model,
                max_tokens=max_tokens,
                temperature=temperature,
                system=system_prompt,
                messages=[{"role": "user", "content": user_prompt}],
            )
            return response.content[0].text.strip()

        except Exception as exc:
            error_msg = str(exc).lower()
            if "context_length_exceeded" in error_msg or "maximum context length" in error_msg:
                self.logger.error("Token limit exceeded, trying with shorter content...")
                if len(user_prompt) > 2000:
                    shortened_prompt = (
                        user_prompt[:2000]
                        + "\n[...noi dung da duoc rut gon do gioi han AI...]"
                    )
                    return self.generate_text(system_prompt, shortened_prompt, max_tokens, temperature)
                return None
            self.logger.error(f"Loi AI generation: {exc}")
            return None


class DocumentTool:
    """Create and summarize documents."""

    @staticmethod
    def create_word_document(title: str, content: str, metadata: Dict[str, Any]) -> bytes:
        doc = Document()
        main_title = doc.add_heading(title, 0)
        main_title.alignment = 1

        if metadata:
            doc.add_heading("Thong tin", level=1)
            for key, value in metadata.items():
                doc.add_paragraph(f"{key}: {value}")

        doc.add_heading("Noi dung", level=1)
        for paragraph in content.split("\n\n"):
            if paragraph.strip():
                doc.add_paragraph(paragraph.strip())

        word_count = len(content.split())
        estimated_time = word_count / 150

        doc.add_heading("Thong ke", level=2)
        stats = doc.add_paragraph()
        stats.add_run(f"So tu: {word_count}\n")
        stats.add_run(f"So ky tu: {len(content)}\n")
        stats.add_run(f"Thoi luong uoc tinh: {estimated_time:.1f} phut\n")
        stats.add_run(f"Thoi gian tao: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    @staticmethod
    def calculate_reading_time(text: str, words_per_minute: int = 150) -> float:
        return len(text.split()) / words_per_minute

    @staticmethod
    def get_text_stats(text: str) -> Dict[str, Any]:
        return {
            "word_count": len(text.split()),
            "char_count": len(text),
            "paragraph_count": len([p for p in text.split("\n\n") if p.strip()]),
            "estimated_time": DocumentTool.calculate_reading_time(text),
        }


class ContentCleanerTool:
    """Helpers to clean extracted content."""

    @staticmethod
    def clean_text(text: str) -> str:
        if not text:
            return ""

        text = re.sub(r"\s+", " ", text)

        contact_patterns = [
            r"Tel:\s*\([^)]+\)[^.]*",
            r"Fax:\s*\([^)]+\)[^.]*",
            r"E-mail:\s*[^\s]+@[^\s]+",
            r"Tru so[^.]*\.",
            r"Dia chi[^.]*\.",
            r"Lien he[^.]*\.",
            r"Hotline[^.]*\.",
            r"Website[^.]*\.",
        ]
        for pattern in contact_patterns:
            text = re.sub(pattern, "", text, flags=re.IGNORECASE)

        ad_keywords = [
            "quang cao",
            "advertisement",
            "sponsored",
            "pr article",
            "bai pr",
            "tai tro",
            "dang ky nhan tin",
        ]
        lines = text.split("\n")
        cleaned_lines = []
        for line in lines:
            line = line.strip()
            if line and not any(keyword in line.lower() for keyword in ad_keywords):
                cleaned_lines.append(line)

        cleaned_text = " ".join(cleaned_lines)
        cleaned_text = re.sub(r"\s+", " ", cleaned_text).strip()
        return cleaned_text

    @staticmethod
    def extract_main_content(soup: BeautifulSoup, content_selectors: list) -> str:
        content_parts = []
        for selector in content_selectors:
            for element in soup.select(selector):
                text = element.get_text(strip=True)
                if text and len(text) > 20:
                    content_parts.append(text)
        if content_parts:
            return ContentCleanerTool.clean_text(" ".join(content_parts))
        return ""

    @staticmethod
    def truncate_content(content: str, max_length: int = 5000) -> str:
        if len(content) <= max_length:
            return content

        truncated = content[:max_length]
        last_sentence = max(truncated.rfind("."), truncated.rfind("!"), truncated.rfind("?"))
        if last_sentence > max_length * 0.8:
            truncated = truncated[: last_sentence + 1]

        return truncated + "\n[...noi dung da duoc rut gon...]"
